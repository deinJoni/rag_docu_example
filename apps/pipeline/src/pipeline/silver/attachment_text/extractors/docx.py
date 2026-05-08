"""DOCX text extraction via python-docx.

Walks paragraphs in document order, then table cells (joined with ' | ' per
row). No attempt at preserving heading levels or styles — embeddings work fine
on flat prose for this corpus's two .docx files.
"""

from __future__ import annotations

import importlib.metadata
from pathlib import Path

from ..models import ExtractionResult

HEURISTIC_VERSION = "v1"


def _version() -> str:
    return f"python-docx@{importlib.metadata.version('python-docx')}+{HEURISTIC_VERSION}"


def extract(path: Path) -> ExtractionResult:
    try:
        from docx import Document  # python-docx exports as module ``docx``

        doc = Document(str(path))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells if c.text]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts)
        return ExtractionResult(
            extractor="python-docx",
            extractor_version=_version(),
            status="ok" if text else "empty",
            text=text,
            char_count=len(text),
        )
    except Exception as err:
        return ExtractionResult(
            extractor="python-docx",
            extractor_version=_version(),
            status="error",
            error_message=str(err),
        )

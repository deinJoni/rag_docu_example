"""One-shot fixture generator for extractor tests.

Run once to (re)generate the binary test fixtures alongside this file:

    uv run python apps/pipeline/tests/fixtures/_generate.py

Outputs:
    sample-text.pdf        — text-bearing PDF (happy path)
    sample-image-only.pdf  — PDF whose only page is a rasterised image
    sample.docx            — small DOCX with a paragraph and a table row
    sample.csv             — header + 2 rows
"""

from __future__ import annotations

from pathlib import Path

HERE = Path(__file__).parent


def make_text_pdf() -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    out = HERE / "sample-text.pdf"
    c = canvas.Canvas(str(out), pagesize=LETTER)
    c.setFont("Helvetica", 14)
    c.drawString(72, 720, "Hello world.")
    c.drawString(72, 700, "Lorem ipsum dolor sit amet, consectetur adipiscing elit.")
    c.drawString(72, 680, "This text is enough to clear the EMPTY_PAGE_THRESHOLD.")
    c.drawString(72, 660, "It also exists on the second page for good measure.")
    c.showPage()
    c.setFont("Helvetica", 12)
    c.drawString(72, 720, "Page 2 — additional content for the multi-page test.")
    c.drawString(72, 700, "Quick brown fox jumps over the lazy dog.")
    c.save()
    print(f"  wrote {out}")


def make_image_only_pdf() -> None:
    """Rasterise a small white PNG and save it as a single-page PDF.

    Pillow's ``save(..., 'PDF')`` produces an image-only PDF with no embedded
    text, which is exactly what we want for the needs_ocr path.
    """
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (612, 792), "white")
    draw = ImageDraw.Draw(img)
    # Draw something visually present but not a vector text object.
    draw.rectangle((100, 100, 500, 400), outline="black", width=4)
    draw.line((100, 100, 500, 400), fill="black", width=2)
    out = HERE / "sample-image-only.pdf"
    img.save(out, "PDF", resolution=72.0)
    print(f"  wrote {out}")


def make_docx() -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello from a docx fixture.")
    doc.add_paragraph("This file exercises the python-docx extractor.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "name"
    table.cell(0, 1).text = "qty"
    table.cell(1, 0).text = "apple"
    table.cell(1, 1).text = "3"
    out = HERE / "sample.docx"
    doc.save(str(out))
    print(f"  wrote {out}")


def make_csv() -> None:
    out = HERE / "sample.csv"
    out.write_text("name,qty\napple,3\npear,5\n", encoding="utf-8")
    print(f"  wrote {out}")


def main() -> None:
    print("Generating extractor test fixtures...")
    make_text_pdf()
    make_image_only_pdf()
    make_docx()
    make_csv()
    print("done.")


if __name__ == "__main__":
    main()
